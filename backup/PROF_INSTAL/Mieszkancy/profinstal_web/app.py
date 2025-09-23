
# -*- coding: utf-8 -*-
from flask import Flask, render_template, request, redirect, url_for, send_file, flash
from datetime import date
import io
from calc import compute_all, compute_audit

app = Flask(__name__)
app.secret_key = "change-me"

DEFAULTS = {
    "bill": 49.00,
    "city": "Kraków",
    "heat_price": 73.69,
    "unit": "GJ",
    "vat": 23.0,
    "month_m3": 7.42,
    "dT": 45,
    "units": 65
}

# Domyślne parametry do audytu (stara i nowa instalacja)
AUDIT_DEFAULTS_OLD = {
    'Q': 20.0,         # kW
    'L': 50.0,         # m
    'd': 32.0,         # mm
    'lambda': 0.035,   # W/mK (np. stara izolacja)
    't_in': 60.0,      # °C
    't_out': 40.0,     # °C
    't_amb': 20.0,     # °C
    'ins_thick': 10.0, # mm
    'czas_pracy': 2000 # h/rok
}
AUDIT_DEFAULTS_NEW = {
    'Q': 20.0,
    'L': 50.0,
    'd': 32.0,
    'lambda': 0.025,   # nowa lepsza izolacja
    't_in': 60.0,
    't_out': 40.0,
    't_amb': 20.0,
    'ins_thick': 30.0, # grubsza izolacja
    'czas_pracy': 2000
}

CITY_PRICES = {
    "Kraków": 73.69,
    "Warszawa": 85.00,
    "Lublin": 65.50,
    "Katowice": 70.00
}

@app.route("/", methods=["GET"])
def index():
    return render_template("index.html", defaults=DEFAULTS, city_prices=CITY_PRICES, audit_defaults_old=AUDIT_DEFAULTS_OLD, audit_defaults_new=AUDIT_DEFAULTS_NEW)
@app.route("/audit", methods=["GET", "POST"])
def audit():
    if request.method == "POST":
        try:
            # Pobierz parametry starej i nowej instalacji z formularza
            params_old = {k: float(request.form.get(f"old_{k}", 0)) for k in AUDIT_DEFAULTS_OLD.keys()}
            params_new = {k: float(request.form.get(f"new_{k}", 0)) for k in AUDIT_DEFAULTS_NEW.keys()}
            
            # Pobierz parametry ceny ciepła (opcjonalne, z domyślnymi wartościami)
            heat_price = float(request.form.get("heat_price", 73.69))
            unit = request.form.get("unit", "GJ")
            vat = float(request.form.get("vat", 23.0))
            
            res = compute_audit(params_old, params_new, heat_price, unit, vat)
            
            # Przygotuj dane do wykresów
            chart_labels = ["Stara instalacja", "Nowa instalacja"]
            heat_losses = [res["Q_loss_old"], res["Q_loss_new"]]
            costs = [res["cost_old"], res["cost_new"]]
            
            return render_template("audit_result.html", 
                                 res=res, 
                                 params_old=params_old, 
                                 params_new=params_new, 
                                 today=date.today().isoformat(),
                                 chart_labels=chart_labels,
                                 heat_losses=heat_losses,
                                 costs=costs)
        except Exception as e:
            flash(f"Błąd danych: {e}")
            return redirect(url_for("audit"))
    return render_template("audit.html", audit_defaults_old=AUDIT_DEFAULTS_OLD, audit_defaults_new=AUDIT_DEFAULTS_NEW)

@app.route("/calc", methods=["POST"])
def calc():
    try:
        bill = float(request.form.get("bill").replace(",", "."))
        city = request.form.get("city") or "Kraków"
        heat_price = float(request.form.get("heat_price").replace(",", "."))
        unit = request.form.get("unit") or "GJ"
        vat = float(request.form.get("vat").replace(",", "."))
        month_m3 = float(request.form.get("month_m3").replace(",", "."))
        dT = float(request.form.get("dT").replace(",", "."))
        units = int(request.form.get("units"))
        res = compute_all(bill, heat_price, unit, vat, month_m3, units, dT)
        return render_template("result.html", res=res, today=date.today().isoformat())
    except Exception as e:
        flash(f"Błąd danych: {e}")
        return redirect(url_for("index"))

# Optional: DOCX/PDF generation stubs (can be expanded later)
@app.route("/export/docx", methods=["POST"])
def export_docx():
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        return "Brak modułu python-docx. Zainstaluj: pip install python-docx", 500

    # parse inputs (quick reuse of calc)
    bill = float(request.form.get("bill").replace(",", "."))
    city = request.form.get("city") or "Kraków"
    heat_price = float(request.form.get("heat_price").replace(",", "."))
    unit = request.form.get("unit") or "GJ"
    vat = float(request.form.get("vat").replace(",", "."))
    month_m3 = float(request.form.get("month_m3").replace(",", "."))
    dT = float(request.form.get("dT").replace(",", "."))
    units = int(request.form.get("units"))
    res = compute_all(bill, heat_price, unit, vat, month_m3, units, dT)

    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("PROF INSTAL — wynik obliczeń (skrót)\n")
    r.bold = True
    doc.add_paragraph(f"Data: {date.today().isoformat()}")
    doc.add_paragraph(f"Rachunek: {res['bill']:.2f} zł/m³; Cena ciepła brutto: {res['price_GJ_brutto']:.2f} zł/GJ; ΔT: {res['dT']:.0f}°C")
    doc.add_paragraph(f"Q_teor: {res['q_per_m3']:.5f} GJ/m³ → koszt_teor: {res['cost_theor']:.2f} zł/m³; η: {res['eta']*100:.1f}%")
    doc.add_paragraph(f"Strata: {res['loss_per_m3']:.2f} zł/m³; Budynek: {res['loss_build_m']:,.2f} zł/m-c; {res['loss_build_y']:,.2f} zł/rok")
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return send_file(bio, as_attachment=True, download_name="PROF_INSTAL_wynik.docx")

@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import mm
    except Exception:
        return "Brak modułu reportlab. Zainstaluj: pip install reportlab", 500

    bill = float(request.form.get("bill").replace(",", "."))
    heat_price = float(request.form.get("heat_price").replace(",", "."))
    unit = request.form.get("unit") or "GJ"
    vat = float(request.form.get("vat").replace(",", "."))
    month_m3 = float(request.form.get("month_m3").replace(",", "."))
    dT = float(request.form.get("dT").replace(",", "."))
    units = int(request.form.get("units"))
    res = compute_all(bill, heat_price, unit, vat, month_m3, units, dT)

    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    w, h = A4
    x, y = 20*mm, h - 20*mm
    lh = 6*mm
    def writeln(text, bold=False, size=10):
        nonlocal y
        c.setFont("Times-Bold" if bold else "Times-Roman", size)
        c.drawString(x, y, text); y -= lh

    writeln("PROF INSTAL — wynik obliczeń (skrót)", bold=True)
    writeln(f"Data: {date.today().isoformat()}")
    writeln(f"Rachunek: {res['bill']:.2f} zł/m³ | Ciepło brutto: {res['price_GJ_brutto']:.2f} zł/GJ | ΔT: {res['dT']:.0f}°C")
    writeln(f"Q_teor: {res['q_per_m3']:.5f} GJ/m³ → koszt_teor: {res['cost_theor']:.2f} zł/m³ | η: {res['eta']*100:.1f}%")
    writeln(f"Strata: {res['loss_per_m3']:.2f} zł/m³ | Budynek: {res['loss_build_m']:,.2f} zł/m-c; {res['loss_build_y']:,.2f} zł/rok")
    c.showPage(); c.save()
    bio.seek(0)
    return send_file(bio, as_attachment=True, download_name="PROF_INSTAL_wynik.pdf", mimetype="application/pdf")

if __name__ == "__main__":
    app.run(debug=True)
