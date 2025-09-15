# -*- coding: utf-8 -*-
"""
PROF INSTAL - Dla Mieszkańców © 2025 Maciej Ślusarczyk
Wszelkie prawa zastrzeżone.

Niniejsze oprogramowanie i jego kod źródłowy są własnością intelektualną 
PROF INSTAL Maciej Ślusarczyk. Kopiowanie, dystrybucja, modyfikacja lub 
wykorzystanie w celach komercyjnych bez pisemnej zgody autora jest zabronione.

Kontakt: prof.instal@example.com
"""

import tkinter as tk
from tkinter import ttk, messagebox, font as tkfont
from math import isfinite
import os, sys, tempfile, subprocess

LOGO_FILE = "profinstal_logo_150x150.png"  # w tym samym folderze

# ========= COPYRIGHT & LICENSING =========
COPYRIGHT_NOTICE = """
© 2025 PROF INSTAL Maciej Ślusarczyk. Wszelkie prawa zastrzeżone.

Niniejsze oprogramowanie i jego kod źródłowy są własnością intelektualną 
PROF INSTAL Maciej Ślusarczyk. Kopiowanie, dystrybucja, modyfikacja lub 
wykorzystanie w celach komercyjnych bez pisemnej zgody autora jest zabronione.

Kontakt: prof.instal@example.com
"""

# ===== moduły opcjonalne =====
DOCX_AVAILABLE = True
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import datetime
except Exception:
    DOCX_AVAILABLE = False

PDF_AVAILABLE = True
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as pdf_canvas
    from reportlab.lib.units import mm
except Exception:
    PDF_AVAILABLE = False

# (logo w GUI)
try:
    from PIL import Image, ImageTk
    PIL_OK = True
except Exception:
    PIL_OK = False

# ===== dane eksperta =====
EXPERT = {
    "name": "mgr inż. Maciej Ślusarczyk",
    "title": "Ekspert HVAC / instalacje sanitarne",
    "lic": "Uprawnienia budowlane bez ograniczeń, nr XXX/XX/XX",
    "chamber": "Członek Małopolskiej OIIB",
    "contact": "kontakt@profinstal.info | +48 123 456 789",
    "company": "PROF INSTAL",
    "city": "Kraków"
}

# ===== paleta (eco dark + delikatny gradient nagłówka) =====
COL_BG      = "#0e1f18"
COL_CARD    = "#173627"
COL_CARD_HL = "#1b3f2d"
COL_TEXT    = "#e9f7ef"
COL_MUTE    = "#b7d8c3"
COL_PRIMARY = "#84c991"
COL_ACCENT  = "#59a868"
COL_WARN    = "#f59e0b"
COL_BAD     = "#dc2626"
COL_GOOD    = "#10b981"

FONT_FAMILY = "Times New Roman"
FONT_HEAD   = (FONT_FAMILY, 13, "bold")
FONT_TEXT   = (FONT_FAMILY, 10)
FONT_BTN    = (FONT_FAMILY, 10, "bold")

# ===== fizyka =====
CP_KJ_PER_KG_K = 4.19  # kJ/(kg·K)

# ===== utils =====
def open_file_crossplatform(path: str):
    try:
        if sys.platform.startswith("win"):
            os.startfile(path)  # type: ignore
        elif sys.platform == "darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
    except Exception as e:
        messagebox.showerror("Nie można otworzyć pliku", f"{path}\n\n{e}")

def logo_path_or_none():
    p = os.path.join(os.path.dirname(__file__), LOGO_FILE)
    return p if os.path.exists(p) else None

# ===== komponent KPI (mini-kart z dużą liczbą + pasek) =====
class KPI(ttk.Frame):
    def __init__(self, parent, title: str, unit: str = "", *, color=COL_PRIMARY):
        super().__init__(parent, style="Card.TFrame", padding=(12, 10))
        self.columnconfigure(0, weight=1)
        self._unit = unit
        self._color = color
        self.title = ttk.Label(self, text=title, style="KPI.Title.TLabel")
        self.title.grid(row=0, column=0, sticky="w")
        self.value = ttk.Label(self, text="—", style="KPI.Value.TLabel")
        self.value.grid(row=1, column=0, sticky="w", pady=(2, 4))
        self.pb = ttk.Progressbar(self, maximum=100)
        self.pb.grid(row=2, column=0, sticky="ew")
    def set(self, val: float, *, percent_for_bar: float | None = None, fmt: str = "{:.2f}"):
        txt = fmt.format(val)
        if self._unit:
            txt += f" {self._unit}"
        self.value.configure(text=txt)
        if percent_for_bar is not None:
            self.pb.configure(value=max(0, min(100, percent_for_bar)))

# ===== GUI =====
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("PROF INSTAL — Dla Mieszkańców © 2025 Maciej Ślusarczyk")
        self.geometry("1200x840")
        self.minsize(1080, 740)
        self.configure(bg=COL_BG)

        # globalny font
        tkfont.nametofont("TkDefaultFont").configure(family=FONT_FAMILY, size=10)
        tkfont.nametofont("TkTextFont").configure(family=FONT_FAMILY)
        tkfont.nametofont("TkHeadingFont").configure(family=FONT_FAMILY)

        self._styles()
        self._ui()

    def _styles(self):
        st = ttk.Style(self)
        st.theme_use("clam")

        # tła
        st.configure("Bg.TFrame", background=COL_BG)
        st.configure("Card.TFrame", background=COL_CARD)
        st.map("Card.TFrame", background=[("active", COL_CARD_HL)])

        # nagłówki i teksty
        st.configure("Title.TLabel", background=COL_CARD, foreground=COL_PRIMARY, font=FONT_HEAD)
        st.configure("Text.TLabel", background=COL_CARD, foreground=COL_TEXT, font=FONT_TEXT)
        st.configure("Info.TLabel", background=COL_BG, foreground=COL_TEXT, font=FONT_TEXT)
        st.configure("Hint.TLabel", background=COL_CARD, foreground=COL_MUTE, font=(FONT_FAMILY, 9))

        # KPI style
        st.configure("KPI.Title.TLabel", background=COL_CARD, foreground=COL_MUTE, font=(FONT_FAMILY, 9))
        st.configure("KPI.Value.TLabel", background=COL_CARD, foreground=COL_TEXT, font=(FONT_FAMILY, 18, "bold"))
        st.configure("TProgressbar", troughcolor="#102a1d", background=COL_ACCENT)

        # pola i przyciski
        st.configure("TEntry", fieldbackground="#f4fff6", foreground="#0f2a17", insertcolor="#0f2a17")
        st.configure("TCombobox", fieldbackground="#f4fff6")
        st.configure("TButton", background=COL_CARD, foreground=COL_PRIMARY, font=FONT_BTN, padding=8)
        st.map("TButton", background=[("active", COL_ACCENT)], foreground=[("active", "#ffffff")])
        st.configure("Accent.TButton", background=COL_PRIMARY, foreground="#0a1e12")
        st.map("Accent.TButton", background=[("active", "#74b882")])
        st.configure("Ghost.TButton", background=COL_CARD, foreground=COL_MUTE)

    # ===== UI =====
    def _ui(self):
        # Główny kontener
        main_container = ttk.Frame(self, style="Bg.TFrame")
        main_container.pack(fill="both", expand=True)

        # Pasek nagłówka z delikatnym gradientem (Canvas)
        header = tk.Canvas(main_container, height=84, highlightthickness=0, bg=COL_BG)
        header.pack(fill="x")
        self._draw_header_gradient(header)
        self._header_contents(header)

        # Główny kontener z przewijaniem
        canvas = tk.Canvas(main_container, bg=COL_BG, highlightthickness=0)
        scrollbar = ttk.Scrollbar(main_container, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, style="Bg.TFrame")

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Bind mouse wheel to canvas
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)

        # Wrap na treść - teraz w scrollable_frame
        wrap = ttk.Frame(scrollable_frame, style="Bg.TFrame", padding=16)
        wrap.pack(fill="both", expand=True)

        # Przyciski w prawym górnym rogu
        top_right_frame = ttk.Frame(wrap, style="Bg.TFrame")
        top_right_frame.pack(fill="x", anchor="e")
        ttk.Button(top_right_frame, text="Dane Mieszkańca", command=self._toggle_sender_form, style="Ghost.TButton").pack(side="right", padx=(0, 8))
        ttk.Button(top_right_frame, text="Dane adresata Spółdzielni / Wspólnoty", command=self._toggle_addressee_form, style="Ghost.TButton").pack(side="right")

        # Grid główny
        grid = ttk.Frame(wrap, style="Bg.TFrame")
        grid.pack(fill="x", pady=(8, 0))

        # Lewa kolumna: dane wejściowe
        self._card_inputs(grid).grid(row=0, column=0, sticky="nsew", padx=(0, 8))

        # Prawa kolumna: nadawca + adresat (początkowo ukryte)
        self.right_column = ttk.Frame(grid, style="Bg.TFrame")
        self.right_column.grid(row=0, column=1, sticky="nsew", padx=(8, 0))
        self.sender_card = self._card_sender(self.right_column)
        self.addressee_card = self._card_addressee(self.right_column)
        
        # Początkowo ukrywamy karty adresowe
        self.sender_visible = False
        self.addressee_visible = False

        grid.columnconfigure(0, weight=1)
        grid.columnconfigure(1, weight=1)

        # Pasek akcji
        self._card_actions(wrap).pack(fill="x", pady=(16, 12))

        # KPI pasek - pierwsze 4 wskaźniki
        self.kpi_bar1 = ttk.Frame(wrap, style="Bg.TFrame")
        self.kpi_bar1.pack(fill="x", pady=(4, 6))
        
        self.kpi_eta   = KPI(self.kpi_bar1, "Sprawność instalacji", "%")
        self.kpi_loss  = KPI(self.kpi_bar1, "Strata na 1 m³", "zł")
        self.kpi_month = KPI(self.kpi_bar1, "Twoja strata / miesiąc", "zł")
        self.kpi_year  = KPI(self.kpi_bar1, "Twoja strata / rok", "zł")
        
        self.kpi_eta.grid(row=0, column=0, sticky="ew", padx=(0, 6))
        self.kpi_loss.grid(row=0, column=1, sticky="ew", padx=6)
        self.kpi_month.grid(row=0, column=2, sticky="ew", padx=6)
        self.kpi_year.grid(row=0, column=3, sticky="ew", padx=(6, 0))
        
        for i in range(4):
            self.kpi_bar1.columnconfigure(i, weight=1)

        # KPI pasek - oszczędności po modernizacji
        self.kpi_bar2 = ttk.Frame(wrap, style="Bg.TFrame")
        self.kpi_bar2.pack(fill="x", pady=(6, 12))
        
        self.kpi_save70_month = KPI(self.kpi_bar2, "Oszczędność 70% / miesiąc", "zł")
        self.kpi_save70_year  = KPI(self.kpi_bar2, "Oszczędność 70% / rok", "zł")
        self.kpi_save80_month = KPI(self.kpi_bar2, "Oszczędność 80% / miesiąc", "zł")
        self.kpi_save80_year  = KPI(self.kpi_bar2, "Oszczędność 80% / rok", "zł")
        
        self.kpi_save70_month.grid(row=0, column=0, sticky="ew", padx=(0, 6))
        self.kpi_save70_year.grid(row=0, column=1, sticky="ew", padx=6)
        self.kpi_save80_month.grid(row=0, column=2, sticky="ew", padx=6)
        self.kpi_save80_year.grid(row=0, column=3, sticky="ew", padx=(6, 0))
        
        for i in range(4):
            self.kpi_bar2.columnconfigure(i, weight=1)

        # Wyniki i wnioski
        self.out_card = self._card_output(wrap)
        self.out_card.pack(fill="both", expand=True, pady=(0, 8))

        ttk.Label(wrap, text="Sprawność η: η = koszt_teor / stawka_rachunkowa (im bliżej 100%, tym lepiej).",
                  style="Info.TLabel").pack(anchor="w", pady=(8, 0))

        # Stopka - w scrollable_frame
        self._create_copyright_footer(scrollable_frame)

    def _toggle_sender_form(self):
        """Pokazuje/ukrywa formularz nadawcy (mieszkańca)"""
        if self.sender_visible:
            # Ukryj formularz nadawcy
            self.sender_card.pack_forget()
            self.sender_visible = False
        else:
            # Pokaż formularz nadawcy
            self.sender_card.pack(fill="x", pady=(0, 8))
            self.sender_visible = True

    def _toggle_addressee_form(self):
        """Pokazuje/ukrywa formularz adresata (spółdzielnia/wspólnota)"""
        if self.addressee_visible:
            # Ukryj formularz adresata
            self.addressee_card.pack_forget()
            self.addressee_visible = False
        else:
            # Pokaż formularz adresata
            if self.sender_visible:
                # Jeśli nadawca jest widoczny, pokaż adresata pod nim
                self.addressee_card.pack(fill="x")
            else:
                # Jeśli nadawca nie jest widoczny, pokaż adresata na górze
                self.addressee_card.pack(fill="x", pady=(0, 8))
            self.addressee_visible = True

    # ===== Nagłówek =====
    def _draw_header_gradient(self, canvas: tk.Canvas):
        # ręcznie malowany pionowy gradient (prosty, bez PIL)
        w = canvas.winfo_reqwidth()
        h = 84
        start = (10, 44, 28)   # ~#0a2c1c
        end   = (18, 62, 39)   # ~#123e27
        for i in range(h):
            t = i / max(h-1,1)
            r = int(start[0] + (end[0]-start[0]) * t)
            g = int(start[1] + (end[1]-start[1]) * t)
            b = int(start[2] + (end[2]-start[2]) * t)
            canvas.create_line(0, i, w, i, fill=f"#{r:02x}{g:02x}{b:02x}")

    def _header_contents(self, canvas: tk.Canvas):
        # Logo + tytuł + podtytuł
        lp = logo_path_or_none()
        if PIL_OK and lp:
            try:
                img = Image.open(lp).resize((64, 64))
                self._logo_img = ImageTk.PhotoImage(img)
                canvas.create_image(24, 42, image=self._logo_img, anchor="w")
            except Exception:
                pass
        canvas.create_text(104, 30, anchor="w", text=f"{EXPERT['company']} — panel mieszkańca CWU",
                           fill=COL_TEXT, font=(FONT_FAMILY, 18, "bold"))
        canvas.create_text(104, 56, anchor="w", text=f"{EXPERT['name']}  |  {EXPERT['title']}",
                           fill=COL_MUTE, font=(FONT_FAMILY, 10))

    # --- karty ---
    def _card_inputs(self, parent):
        c = ttk.Frame(parent, style="Card.TFrame", padding=16)
        ttk.Label(c, text="Dane techniczno-kosztowe - przepisz z rachunku", style="Title.TLabel").grid(row=0, column=0, columnspan=4, sticky="w", pady=(0, 10))

        self._lbl(c, "Opłata z rachunku za podgrzanie wody [zł/m³]:", 1); self.e_bill = self._entry(c, "49.00", 1)
        
        # Miasto i cena ciepła
        self._lbl(c, "Miasto:", 2)
        self.city_var = tk.StringVar(value="Kraków")
        city_combo = ttk.Combobox(c, textvariable=self.city_var, values=["Kraków", "Warszawa", "Lublin", "Katowice"], width=12, state="readonly")
        city_combo.grid(row=2, column=1, sticky="w", padx=(8, 20))
        city_combo.bind("<<ComboboxSelected>>", self._on_city_change)
        
        self._lbl(c, "Cena ciepła MPEC/PEC — netto:", 3); self.e_heat_price = self._entry(c, "73.69", 3)
        self.unit_var = tk.StringVar(value="GJ")
        ttk.Radiobutton(c, text="zł/GJ", value="GJ", variable=self.unit_var).grid(row=3, column=2, sticky="w")
        ttk.Radiobutton(c, text="zł/MJ", value="MJ", variable=self.unit_var).grid(row=3, column=3, sticky="w")

        self._lbl(c, "VAT dla ciepła:", 4)
        ttk.Label(c, text="23%", style="Text.TLabel").grid(row=4, column=1, sticky="w", padx=(8, 20))

        self._lbl(c, "Zużycie miesięczne [m³]:", 5); self.e_month = self._entry(c, "7.42", 5)
        
        self._lbl(c, "ΔT podgrzewu [°C]:", 6)
        ttk.Label(c, text="45", style="Text.TLabel").grid(row=6, column=1, sticky="w", padx=(8, 20))
        
        self._lbl(c, "Liczba podobnych mieszkań:", 7); self.e_units = self._entry(c, "65", 7)

        ttk.Label(c,
                  text="Cele porównawcze: 70% i 80% sprawności (izolacje, regulacja cyrkulacji, przeglądy węzła).",
                  style="Hint.TLabel").grid(row=8, column=0, columnspan=4, sticky="w", pady=(12, 0))
        
        for i in range(9): c.rowconfigure(i, weight=0)
        for j in range(4): c.columnconfigure(j, weight=1)
        return c

    def _on_city_change(self, event=None):
        """Aktualizuje cenę ciepła na podstawie wybranego miasta"""
        city_prices = {
            "Kraków": "73.69",
            "Warszawa": "85.00",
            "Lublin": "65.50",
            "Katowice": "70.00"
        }
        self.e_heat_price.delete(0, tk.END)
        self.e_heat_price.insert(0, city_prices.get(self.city_var.get(), "73.69"))

    def _lbl(self, parent, text, r):
        ttk.Label(parent, text=text, style="Text.TLabel").grid(row=r, column=0, sticky="w", pady=(6, 4))

    def _entry(self, parent, default, r):
        e = ttk.Entry(parent, width=12); e.insert(0, default)
        e.grid(row=r, column=1, sticky="w", padx=(8, 20), pady=(6, 4))
        return e

    def _card_sender(self, parent):
        c = ttk.Frame(parent, style="Card.TFrame", padding=16)
        ttk.Label(c, text="Nadawca (mieszkaniec)", style="Title.TLabel").grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 10))
        self.e_res_name = self._mk_pair(c, "Imię i nazwisko:", "Jan Kowalski", 1)
        self.e_res_addr = self._mk_pair(c, "Adres lokalu:", "ul. Przykładowa 12/34, 30-000 Kraków", 2)
        self.e_res_email = self._mk_pair(c, "E-mail:", "jan.kowalski@example.com", 3)
        self.e_res_phone = self._mk_pair(c, "Telefon:", "+48 600 000 000", 4)
        self._mk_pair(c, "Nr mieszkania/identyfikator:", "lok. 34 / kl. B", 5, varname="e_flat_id")
        return c

    def _card_addressee(self, parent):
        c = ttk.Frame(parent, style="Card.TFrame", padding=16)
        ttk.Label(c, text="Adresat (spółdzielnia / wspólnota)", style="Title.TLabel").grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 10))
        self._mk_pair(c, "Nazwa:", "Spółdzielnia Mieszkaniowa XYZ", 1, varname="e_add_name")
        self._mk_pair(c, "Adres:", "ul. Zarządcza 1, 30-000 Kraków", 2, varname="e_add_addr")
        self._mk_pair(c, "E-mail:", "biuro@sm-xyz.pl", 3, varname="e_add_email")
        self._mk_pair(c, "NIP (opcjonalnie):", "", 4, varname="e_add_nip")
        return c

    def _mk_pair(self, parent, label, default, r, varname=None):
        ttk.Label(parent, text=label, style="Text.TLabel").grid(row=r, column=0, sticky="w", pady=(6, 4))
        e = ttk.Entry(parent, width=28); e.insert(0, default)
        e.grid(row=r, column=1, sticky="ew", padx=(8, 0), pady=(6, 4))
        parent.columnconfigure(1, weight=1)
        if varname:
            setattr(self, varname, e)
        return e

    def _card_actions(self, parent):
        c = ttk.Frame(parent, style="Card.TFrame", padding=14)
        ttk.Label(c, text="Akcje", style="Title.TLabel").grid(row=0, column=0, sticky="w", padx=(4, 16), pady=(0, 8))
        ttk.Button(c, text="Oblicz straty i oszczędności", style="Accent.TButton", command=self._calc_all).grid(row=0, column=1, padx=4)
        ttk.Button(c, text="Pismo reklamacyjne (DOCX)", command=self._gen_resident_letter_docx, style="Ghost.TButton").grid(row=0, column=2, padx=4)
        ttk.Button(c, text="Pismo reklamacyjne (PDF)",  command=self._gen_resident_letter_pdf, style="Ghost.TButton").grid(row=0, column=3, padx=4)
        ttk.Button(c, text="Opinia eksperta (DOCX)", command=self._gen_opinion_docx, style="Ghost.TButton").grid(row=0, column=4, padx=4)
        ttk.Button(c, text="Opinia eksperta (PDF)",  command=self._gen_opinion_pdf, style="Ghost.TButton").grid(row=0, column=5, padx=4)
        for i in range(6): c.columnconfigure(i, weight=1)
        return c

    def _card_output(self, parent):
        c = ttk.Frame(parent, style="Card.TFrame", padding=16)
        ttk.Label(c, text="Wyniki i wzory", style="Title.TLabel").grid(row=0, column=0, sticky="w", pady=(0, 8))
        self.txt = tk.Text(c, height=16, relief="flat", bg="#1a3e25", fg=COL_TEXT, insertbackground=COL_TEXT, wrap="word")
        self.txt.grid(row=1, column=0, sticky="nsew", pady=(0, 0))
        c.rowconfigure(1, weight=1)
        c.columnconfigure(0, weight=1)
        return c

    # ===== obliczenia =====
    def _read_float(self, entry):
        v = float(entry.get().replace(",", "."))
        if not isfinite(v): raise ValueError
        return v

    def _price_GJ_brutto(self, price_net, unit, vat_percent):
        """
        Model przeliczania cen energii - autorski algorytm PROF INSTAL
        © 2025 Maciej Ślusarczyk - Własność intelektualna
        """
        per_GJ_net = price_net * 1000.0 if unit == "MJ" else price_net
        return per_GJ_net * (1.0 + float(vat_percent)/100.0)

    def _Q_GJ_per_m3(self, dT):
        """
        Model fizyczny podgrzewu CWU - autorski wzór PROF INSTAL
        © 2025 Maciej Ślusarczyk - Własność intelektualna
        
        Q = m·c·ΔT; m=1000 kg; c=4,19 kJ/(kg·K); GJ = kJ/1e6
        """
        return (1000.0 * CP_KJ_PER_KG_K * dT) / 1_000_000.0

    def _calc_all(self):
        try:
            bill      = self._read_float(self.e_bill)
            heat_price= self._read_float(self.e_heat_price)
            unit      = self.unit_var.get()
            vat       = "23"  # Stała wartość VAT 23%
            month_m3  = self._read_float(self.e_month)
            units     = int(self.e_units.get())
            dT        = 45.0  # Stała wartość ΔT = 45°C

            price_GJ_brutto = self._price_GJ_brutto(heat_price, unit, vat)
            q_per_m3        = self._Q_GJ_per_m3(dT)
            cost_theor      = q_per_m3 * price_GJ_brutto
            eta             = max(min(cost_theor / bill, 1.0), 0.0)

            loss_per_m3   = bill - cost_theor
            loss_flat_m   = loss_per_m3 * month_m3
            loss_build_m  = loss_flat_m * units
            loss_build_y  = loss_build_m * 12.0

            def cost_at_eff(eff): return (q_per_m3 / eff) * price_GJ_brutto
            cost70 = cost_at_eff(0.70); cost80 = cost_at_eff(0.80)
            save70_m3 = max(bill - cost70, 0.0); save80_m3 = max(bill - cost80, 0.0)
            save70_flat_m = save70_m3 * month_m3; save80_flat_m = save80_m3 * month_m3
            save70_build_m = save70_flat_m * units; save80_build_m = save80_flat_m * units
            save70_build_y = save70_build_m * 12.0; save80_build_y = save80_build_m * 12.0

            self._print_results(bill, heat_price, unit, vat, dT, q_per_m3, price_GJ_brutto,
                                cost_theor, eta, loss_per_m3, loss_flat_m, loss_build_m, loss_build_y,
                                cost70, cost80, save70_m3, save80_m3, save70_flat_m, save80_flat_m,
                                save70_build_m, save80_build_m, save70_build_y, save80_build_y, month_m3, units)

            # KPI aktualizacja
            self.kpi_eta.set(eta*100.0, percent_for_bar=eta*100.0, fmt="{:.1f}")
            self.kpi_loss.set(max(loss_per_m3, 0.0), percent_for_bar=None, fmt="{:.2f}")
            self.kpi_month.set(max(loss_flat_m, 0.0), percent_for_bar=None, fmt="{:.2f}")
            self.kpi_year.set(max(loss_flat_m * 12, 0.0), percent_for_bar=None, fmt="{:.0f}")
            
            # KPI oszczędności po modernizacji
            self.kpi_save70_month.set(max(save70_flat_m, 0.0), percent_for_bar=None, fmt="{:.2f}")
            self.kpi_save70_year.set(max(save70_flat_m * 12, 0.0), percent_for_bar=None, fmt="{:.0f}")
            self.kpi_save80_month.set(max(save80_flat_m, 0.0), percent_for_bar=None, fmt="{:.2f}")
            self.kpi_save80_year.set(max(save80_flat_m * 12, 0.0), percent_for_bar=None, fmt="{:.0f}")

            # bufor ostatnich danych
            self._last = dict(
                bill=bill, heat_price=heat_price, unit=unit, vat=vat, dT=dT, month_m3=month_m3, units=units,
                price_GJ_brutto=price_GJ_brutto, q_per_m3=q_per_m3, cost_theor=cost_theor, eta=eta,
                loss_per_m3=loss_per_m3, loss_flat_m=loss_flat_m, loss_build_m=loss_build_m, loss_build_y=loss_build_y,
                cost70=cost70, cost80=cost80, save70_m3=save70_m3, save80_m3=save80_m3,
                save70_flat_m=save70_flat_m, save80_flat_m=save80_flat_m,
                save70_build_m=save70_build_m, save80_build_m=save80_build_m,
                save70_build_y=save70_build_y, save80_build_y=save80_build_y
            )

        except Exception as e:
            messagebox.showerror("Błąd danych", f"Sprawdź wprowadzone wartości.\n\nSzczegóły: {e}")

    def _print_results(self, bill, heat_price, unit, vat, dT, q_per_m3, price_GJ_brutto,
                       cost_theor, eta, loss_m3, loss_flat_m, loss_build_m, loss_build_y,
                       cost70, cost80, save70_m3, save80_m3, save70_flat_m, save80_flat_m,
                       save70_build_m, save80_build_m, save70_build_y, save80_build_y, month_m3, units):
        self.txt.delete("1.0", "end")
        W = self.txt
        def ln(s=""): W.insert("end", s + "\n")
        ln("PROF INSTAL - ANALIZA CWU © 2025 Maciej Ślusarczyk")
        ln("Wszelkie prawa zastrzeżone - Własność intelektualna")
        ln("=" * 60)
        ln()
        ln("— WZORY AUTORSKIE —")
        ln("Q = m·c·ΔT;  m=1000 kg;  c=4,19 kJ/(kg·K)")
        ln("Q[GJ/m³] = (1000·4,19·ΔT) / 1 000 000")
        ln("koszt_teor [zł/m³] = Q · cena_ciepła_brutto [zł/GJ]")
        ln("η = koszt_teor / stawka_rachunkowa")
        ln()
        ln("— WYNIKI OBLICZEŃ —")
        ln(f"Rachunek: {bill:.2f} zł/m³ | Ciepło brutto: {price_GJ_brutto:.2f} zł/GJ | ΔT: {dT:.0f}°C")
        ln(f"Q_teor: {q_per_m3:.5f} GJ/m³  →  koszt_teor: {cost_theor:.2f} zł/m³ | η: {eta*100:,.1f}%")
        ln(f"Strata: {loss_m3:.2f} zł/m³")
        ln(f"Twoja strata /miesiąc: {loss_flat_m:.2f} zł (przy zużyciu {month_m3:.2f} m³/m-c)")
        ln(f"Twoja strata /rok: {loss_flat_m * 12:.2f} zł")
        ln(f"Budynek łącznie: {loss_build_m:,.2f} zł/m-c; {loss_build_y:,.2f} zł/rok")
        ln()
        ln("— SCENARIUSZE MODERNIZACJI —")
        ln(f"70%: koszt {cost70:.2f} zł/m³ → oszcz. {save70_m3:.2f} zł/m³ | budynek {save70_build_m:,.2f} zł/m-c; {save70_build_y:,.2f} zł/rok")
        ln(f"80%: koszt {cost80:.2f} zł/m³ → oszcz. {save80_m3:.2f} zł/m³ | budynek {save80_build_m:,.2f} zł/m-c; {save80_build_y:,.2f} zł/rok")
        ln()
        ln("© 2025 PROF INSTAL Maciej Ślusarczyk - Model obliczeniowy zastrzeżony")

    # ===== DOCX helpers =====
    def _style_docx_times(self, doc: Document):
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        style.font.size = Pt(11)

    def _add_docx_header(self, doc: Document):
        # Copyright na początku dokumentu
        copyright_para = doc.add_paragraph()
        copyright_run = copyright_para.add_run(COPYRIGHT_NOTICE)
        copyright_run.font.size = Pt(8)
        copyright_run.italic = True
        doc.add_paragraph("")

        table = doc.add_table(rows=1, cols=2)
        table.columns[0].width = Inches(5.5)
        table.columns[1].width = Inches(1.5)
        cell_l, cell_r = table.rows[0].cells
        p = cell_l.paragraphs[0]
        p.add_run(f"{EXPERT['company']} — {EXPERT['city']}\n").bold = True
        p.add_run(f"{EXPERT['name']} | {EXPERT['title']}\n{EXPERT['lic']}\n{EXPERT['chamber']}\n{EXPERT['contact']}")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp = logo_path_or_none()
        if lp:
            try:
                cell_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell_r.paragraphs[0].add_run().add_picture(lp, width=Inches(1.0))
            except Exception:
                pass
        doc.add_paragraph("")

    def _legal_basis_docx(self, doc: Document):
        doc.add_heading("Podstawy prawne i normatywne", level=1)
        doc.add_paragraph("• Ustawa o spółdzielniach mieszkaniowych – art. 4 ust. 2 (rzeczywiste koszty).")
        doc.add_paragraph("• Prawo budowlane – art. 62 (należyty stan techniczny; przeglądy okresowe).")
        doc.add_paragraph("• Prawo energetyczne – art. 5 (należyta staranność i efektywność usług energetycznych).")
        doc.add_paragraph("• Warunki Techniczne (rozp. MI) m.in. §120, §134 – instalacje wodociągowe/CWU; ograniczanie strat; cyrkulacja.")
        doc.add_paragraph("• Taryfy ciepła zatwierdzane przez Prezesa URE – ceny energii (zł/GJ) i opłaty dystrybucyjne.")

    def _signature_block_docx(self, doc: Document):
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(EXPERT["name"] + "\n").bold = True
        p.add_run(EXPERT["title"] + "\n")
        p.add_run(EXPERT["lic"] + "\n")
        p.add_run(EXPERT["chamber"] + "\n")
        p.add_run(EXPERT["contact"])
        doc.add_paragraph("")
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(f"Oprogramowanie: PROF INSTAL - Dla Mieszkańców © 2025 Maciej Ślusarczyk")
        footer_run.font.size = Pt(8)
        footer_run.bold = True

    # ===== PISMO REKLAMACYJNE (DOCX / PDF) =====
    def _collect_parties(self):
        return dict(
            r_name=getattr(self, 'e_res_name', ttk.Entry()).get().strip() if hasattr(self, 'e_res_name') else "Jan Kowalski",
            r_addr=getattr(self, 'e_res_addr', ttk.Entry()).get().strip() if hasattr(self, 'e_res_addr') else "",
            r_email=getattr(self, 'e_res_email', ttk.Entry()).get().strip() if hasattr(self, 'e_res_email') else "",
            r_phone=getattr(self, 'e_res_phone', ttk.Entry()).get().strip() if hasattr(self, 'e_res_phone') else "",
            flatid=getattr(self, 'e_flat_id', ttk.Entry()).get().strip() if hasattr(self, 'e_flat_id') else "",
            a_name=getattr(self, 'e_add_name', ttk.Entry()).get().strip() if hasattr(self, 'e_add_name') else "Spółdzielnia Mieszkaniowa XYZ",
            a_addr=getattr(self, 'e_add_addr', ttk.Entry()).get().strip() if hasattr(self, 'e_add_addr') else "",
            a_email=getattr(self, 'e_add_email', ttk.Entry()).get().strip() if hasattr(self, 'e_add_email') else "",
            a_nip=getattr(self, 'e_add_nip', ttk.Entry()).get().strip() if hasattr(self, 'e_add_nip') else ""
        )

    def _gen_resident_letter_docx(self):
        if not DOCX_AVAILABLE:
            messagebox.showwarning("Brak modułu", "Zainstaluj: py -m pip install python-docx")
            return
        try:
            if not hasattr(self, "_last"): self._calc_all()
            L = self._last; P = self._collect_parties()

            doc = Document(); self._style_docx_times(doc)

            copyright_para = doc.add_paragraph()
            copyright_run = copyright_para.add_run(COPYRIGHT_NOTICE)
            copyright_run.font.size = Pt(8)
            copyright_run.italic = True
            doc.add_paragraph("")

            doc.add_paragraph(P["r_name"]); doc.add_paragraph(P["r_addr"])
            if P["r_email"] or P["r_phone"]:
                doc.add_paragraph(f"{P['r_email']} | {P['r_phone']}")
            doc.add_paragraph("")
            doc.add_paragraph(P["a_name"]); doc.add_paragraph(P["a_addr"])
            if P["a_email"]:
                doc.add_paragraph(P["a_email"])
            if P["a_nip"]:
                doc.add_paragraph(f"NIP: {P['a_nip']}")
            doc.add_paragraph("")
            doc.add_paragraph(f"Data: {datetime.datetime.today().date()}")
            doc.add_paragraph("")
            doc.add_heading("Reklamacja dotycząca zawyżonych kosztów podgrzania ciepłej wody użytkowej", 0)

            doc.add_paragraph(
                f"Jako mieszkaniec ({P['flatid']}) składam reklamację w zakresie kosztów podgrzewu CWU. "
                f"Na podstawie niezależnej opinii eksperckiej PROF INSTAL wyliczona sprawność instalacji wspólnej "
                f"wynosi około {L['eta']*100:.1f}%. Realny koszt podgrzania 1 m³ powinien wynosić ok. "
                f"{L['cost_theor']:.2f} zł/m³, podczas gdy na rachunku płacę {L['bill']:.2f} zł/m³."
            )
            doc.add_paragraph(
                f"Różnica (moja strata) to {L['loss_per_m3']:.2f} zł na każdym m³. Przy miesięcznym zużyciu "
                f"{L['month_m3']:.2f} m³ daje to {L['loss_flat_m']:.2f} zł/miesiąc. To są pieniądze, które wypływają, "
                f"bo instalacja po stronie zarządcy nie działa wystarczająco sprawnie."
            )
            doc.add_paragraph(
                "Żądam uwzględnienia reklamacji, przedstawienia planu naprawczego (izolacje, regulacja/równoważenie cyrkulacji, "
                "przegląd i korekta nastaw węzła) oraz korekty rozliczeń tak, aby nie obciążać mieszkańców kosztami strat. "
                "Proszę o pisemną odpowiedź w terminie 14 dni."
            )
            doc.add_paragraph("")
            doc.add_paragraph("Z poważaniem,")
            doc.add_paragraph(P["r_name"])

            doc.add_paragraph("")
            doc.add_paragraph("Na żądanie udostępnię pełną opinię eksperta PROF INSTAL z obliczeniami.")
            doc.add_paragraph("")
            footer_para = doc.add_paragraph()
            footer_run = footer_para.add_run(f"Analiza wykonana przy użyciu: PROF INSTAL - Dla Mieszkańców © 2025 Maciej Ślusarczyk")
            footer_run.font.size = Pt(8)
            footer_run.italic = True

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            path = tmp.name; tmp.close()
            doc.save(path); open_file_crossplatform(path)
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się utworzyć pisma (DOCX).\n{e}")

    def _gen_resident_letter_pdf(self):
        if not PDF_AVAILABLE:
            messagebox.showwarning("Brak modułu", "Zainstaluj: py -m pip install reportlab")
            return
        try:
            if not hasattr(self, "_last"): self._calc_all()
            L = self._last; P = self._collect_parties()

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            path = tmp.name; tmp.close()

            c = pdf_canvas.Canvas(path, pagesize=A4)
            w, h = A4
            x, y = 20*mm, h - 20*mm
            lh = 6*mm

            def writeln(text, bold=False, size=10):
                nonlocal y
                c.setFont("Times-Bold" if bold else "Times-Roman", size)
                c.drawString(x, y, text)
                y -= lh

            lp = logo_path_or_none()
            if lp:
                try: c.drawImage(lp, w-40*mm, h-30*mm, width=20*mm, height=20*mm, preserveAspectRatio=True, mask='auto')
                except Exception: pass

            writeln("© 2025 PROF INSTAL Maciej Ślusarczyk - Wszelkie prawa zastrzeżone", size=8)
            y -= 2*mm

            writeln(f"{EXPERT['company']} — {EXPERT['city']}", bold=True, size=12)
            writeln(f"{EXPERT['name']} | {EXPERT['title']}", size=9)
            writeln(f"{EXPERT['lic']} | {EXPERT['chamber']}", size=9)
            writeln(EXPERT["contact"], size=9); y -= 4*mm

            writeln("Opinia techniczno-finansowa — CWU (skrót)", bold=True); y -= 2*mm
            writeln("Wzory autorskie PROF INSTAL © 2025 Maciej Ślusarczyk:", bold=True, size=9)
            writeln("Q = m·c·ΔT;  Q[GJ/m³]=(1000·4,19·ΔT)/1e6;  koszt_teor = Q·cena_ciepła_brutto;  η = koszt_teor/rachunek")
            writeln(f"Rachunek: {L['bill']:.2f} zł/m³ | Ciepło brutto: {L['price_GJ_brutto']:.2f} zł/GJ | ΔT: {L['dT']:.0f}°C")
            writeln(f"Q_teor: {L['q_per_m3']:.5f} GJ/m³ → koszt_teor: {L['cost_theor']:.2f} zł/m³ | η: {L['eta']*100:.1f}%")
            writeln(f"Strata: {L['loss_per_m3']:.2f} zł/m³ | Budynek: {L['loss_build_m']:,.2f} zł/m-c; {L['loss_build_y']:,.2f} zł/rok"); y -= 2*mm
            writeln("Scenariusze modernizacji:", bold=True)
            writeln(f"70% → koszt {L['cost70']:.2f} zł/m³ | oszcz. {L['save70_m3']:.2f} zł/m³")
            writeln(f"80% → koszt {L['cost80']:.2f} zł/m³ | oszcz. {L['save80_m3']:.2f} zł/m³")
            y -= 2*mm
            writeln("Oprogramowanie: PROF INSTAL - Dla Mieszkańców © 2025 Maciej Ślusarczyk", size=8)

            c.showPage(); c.save(); open_file_crossplatform(path)
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się utworzyć opinii (PDF).\n{e}")

    # ===== OPINIA EKSPERTA (DOCX / PDF) =====
    def _gen_opinion_docx(self):
        if not DOCX_AVAILABLE:
            messagebox.showwarning("Brak modułu", "Zainstaluj: py -m pip install python-docx")
            return
        try:
            if not hasattr(self, "_last"): self._calc_all()
            L = self._last

            doc = Document(); self._style_docx_times(doc); self._add_docx_header(doc)
            doc.add_heading("Opinia techniczno-finansowa — ciepła woda użytkowa (CWU)", 0)
            doc.add_paragraph(f"Data: {datetime.datetime.today().date()}")

            doc.add_heading("1. Metodyka i wzory obliczeniowe", level=1)
            doc.add_paragraph("Obliczenia wykonano według zasad bilansu cieplnego (Q = m·c·ΔT) opisanych w normach i literaturze technicznej, z autorską metodyką PROF INSTAL w zakresie interpretacji sprawności, analizy strat i symulacji scenariuszy modernizacji.")
            doc.add_paragraph("Q = m·c·ΔT;  m = 1000 kg;  c = 4,19 kJ/(kg·K).")
            doc.add_paragraph("Q[GJ/m³] = (1000·4,19·ΔT) / 1 000 000;  koszt_teor [zł/m³] = Q · cena_ciepła_brutto [zł/GJ].")
            doc.add_paragraph("Sprawność: η = koszt_teor / stawka_rachunkowa.")

            method_para = doc.add_paragraph()
            method_run = method_para.add_run("© 2025 PROF INSTAL Maciej Ślusarczyk - Model obliczeniowy objęty prawami autorskimi")
            method_run.font.size = Pt(9)
            method_run.italic = True

            doc.add_heading("2. Dane wejściowe i wyniki", level=1)
            u = "zł/GJ" if L['unit']=="GJ" else "zł/MJ"
            doc.add_paragraph(f"Stawka rachunkowa: {L['bill']:.2f} zł/m³")
            doc.add_paragraph(f"Cena ciepła (netto): {L['heat_price']:.4f} {u}; VAT: {L['vat']}%; Ciepło brutto: {L['price_GJ_brutto']:.2f} zł/GJ")
            doc.add_paragraph(f"ΔT: {L['dT']:.0f}°C  →  Q_teor = {L['q_per_m3']:.5f} GJ/m³;  koszt_teor = {L['cost_theor']:.2f} zł/m³")
            doc.add_paragraph(f"Sprawność wyliczona: {L['eta']*100:.1f}%")
            doc.add_paragraph(f"Strata na 1 m³: {L['loss_per_m3']:.2f} zł;  Strata budynku: {L['loss_build_m']:,.2f} zł/m-c; {L['loss_build_y']:,.2f} zł/rok.")

            # Dodanie szczegółowych obliczeń krok po kroku
            doc.add_heading("2.1. Szczegółowe obliczenia krok po kroku", level=2)
            
            doc.add_paragraph("Obliczenie zapotrzebowania na energię:")
            doc.add_paragraph(f"Q = 1000 · 4,19 · {L['dT']:.0f} / 1 000 000 = {L['q_per_m3']:.5f} GJ/m³")
            
            doc.add_paragraph("Cena brutto energii cieplnej:")
            doc.add_paragraph(f"{L['heat_price']:.2f} · 1,{L['vat']} = {L['price_GJ_brutto']:.2f} zł/GJ")
            
            doc.add_paragraph("Koszt teoretyczny podgrzania 1 m³:")
            doc.add_paragraph(f"{L['q_per_m3']:.5f} · {L['price_GJ_brutto']:.2f} = {L['cost_theor']:.2f} zł/m³")
            
            doc.add_paragraph('Sprawność (definiowana jako "efektywność rozliczeniowa"):')
            doc.add_paragraph(f"η = {L['cost_theor']:.2f} / {L['bill']:.2f} = {L['eta']*100:.1f}%")
            
            doc.add_paragraph("Strata na 1 m³:")
            doc.add_paragraph(f"{L['bill']:.2f} - {L['cost_theor']:.2f} = {L['loss_per_m3']:.2f} zł")
            
            # Zużycie budynku
            total_consumption_month = L['loss_build_m'] / L['loss_per_m3'] if L['loss_per_m3'] > 0 else 0
            total_consumption_year = total_consumption_month * 12
            doc.add_paragraph(f"Z podanych danych wynika całkowite zużycie budynku:")
            doc.add_paragraph(f"~{total_consumption_month:.1f} m³/miesiąc oraz ~{total_consumption_year:.1f} m³/rok")

            doc.add_heading("3. Scenariusze po modernizacji", level=1)
            doc.add_paragraph("Scenariusze liczone według wzoru: stawka = koszt_teor / η")
            
            doc.add_paragraph(f"70% sprawności:")
            cost70_calc = L['cost_theor'] / 0.70
            save70_calc = L['bill'] - cost70_calc
            doc.add_paragraph(f"{L['cost_theor']:.2f} / 0,7 = {cost70_calc:.2f} zł/m³ → oszczędność {L['bill']:.2f} - {cost70_calc:.2f} = {save70_calc:.2f} zł/m³")
            
            doc.add_paragraph(f"80% sprawności:")
            cost80_calc = L['cost_theor'] / 0.80
            save80_calc = L['bill'] - cost80_calc
            doc.add_paragraph(f"{L['cost_theor']:.2f} / 0,8 = {cost80_calc:.2f} zł/m³ → oszczędność {L['bill']:.2f} - {cost80_calc:.2f} = {save80_calc:.2f} zł/m³")
            
            doc.add_paragraph(f"Oszczędność budynku (mies.): 70% → {L['save70_build_m']:,.2f} zł; 80% → {L['save80_build_m']:,.2f} zł.")
            doc.add_paragraph(f"Oszczędność budynku (rok):   70% → {L['save70_build_y']:,.2f} zł; 80% → {L['save80_build_y']:,.2f} zł.")

            doc.add_heading("4. Uwagi metodyczne", level=1)
            doc.add_paragraph("• Sprawność rozliczeniowa obejmuje straty przesyłu, magazynowania, cyrkulacji, przegrzewy anty-Legionella itd., a nie tylko sprawność wymiennika.")
            doc.add_paragraph("• Wartość c = 4,19 kJ/(kg·K) jest przyjęta dla zakresu temperatur 10-60°C. W precyzyjnych obliczeniach można uwzględnić zmienność c i gęstości ρ≈998 kg/m³.")
            doc.add_paragraph("• Stawka rachunkowa i cena ciepła brutto odnoszą się do tej samej bazy podatkowej.")
            doc.add_paragraph("• Wyniki mogą podlegać niewielkim wahaniom sezonowym ze względu na zmienność temperatury wody zimnej.")

            doc.add_heading("5. Podstawy prawne i normatywne", level=1)
            self._legal_basis_docx(doc)

            doc.add_heading("6. Ocena ekspercka i zalecenia", level=1)
            eta_pct = L['eta']*100
            if eta_pct >= 70:
                ocena = ("Sprawność oceniona jako dobra. Rekomenduję utrzymanie parametrów, okresowe równoważenie cyrkulacji, "
                         "monitoring temperatur (zasilanie/powrót CWU) i audyt co 12 miesięcy.")
                # Dodatkowa notatka dla sprawności w okolicy 70%
                if 70 <= eta_pct <= 75:
                    doc.add_paragraph("")
                    note_para = doc.add_paragraph()
                    note_run = note_para.add_run("UWAGA: Sprawność na granicy normy technicznej.")
                    note_run.bold = True
                    doc.add_paragraph("Mimo że instalacja osiąga akceptowalny poziom 70%, warto rozważyć niewielkie usprawnienia: "
                                    "optymalizację temperatur cyrkulacji, częściowe docieplenia najbardziej narażonych odcinków "
                                    "oraz poprawę regulacji automatyki. Działania te mogą podnieść sprawność do 75-80% "
                                    "przy relatywnie niskich nakładach finansowych.")
            elif 50 <= eta_pct < 70:
                ocena = ("Sprawność umiarkowana — realny potencjał poprawy 10–30%. Priorytety: docieplenia, zawory termostatyczne i równoważenie cyrkulacji, "
                         "korekta nastaw węzła, praca pomp wg temperatury powrotu/harmonogramu.")
            else:
                ocena = ("Sprawność niska — nadmierne straty energii. Zalecane pilne działania: pełne izolacje przewodów, "
                         "równoważenie i sterowanie cyrkulacji, przegląd wymienników i automatyki, eliminacja przegrzewów.")
            doc.add_paragraph(ocena)

            doc.add_heading("7. Metryka i podpis", level=1)
            self._signature_block_docx(doc)

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            path = tmp.name; tmp.close()
            doc.save(path); open_file_crossplatform(path)
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się utworzyć opinii (DOCX).\n{e}")

    def _gen_opinion_pdf(self):
        if not PDF_AVAILABLE:
            messagebox.showwarning("Brak modułu", "Zainstaluj: py -m pip install reportlab")
            return
        try:
            if not hasattr(self, "_last"): self._calc_all()
            L = self._last

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            path = tmp.name; tmp.close()

            c = pdf_canvas.Canvas(path, pagesize=A4)
            w, h = A4
            x, y = 20*mm, h - 20*mm
            lh = 6*mm

            def writeln(text, bold=False, size=10):
                nonlocal y
                c.setFont("Times-Bold" if bold else "Times-Roman", size)
                c.drawString(x, y, text)
                y -= lh

            lp = logo_path_or_none()
            if lp:
                try: c.drawImage(lp, w-40*mm, h-30*mm, width=20*mm, height=20*mm, preserveAspectRatio=True, mask='auto')
                except Exception: pass

            writeln("© 2025 PROF INSTAL Maciej Ślusarczyk - Wszelkie prawa zastrzeżone", size=8)
            y -= 2*mm

            writeln(f"{EXPERT['company']} — {EXPERT['city']}", bold=True, size=12)
            writeln(f"{EXPERT['name']} | {EXPERT['title']}", size=9)
            writeln(f"{EXPERT['lic']} | {EXPERT['chamber']}", size=9)
            writeln(EXPERT["contact"], size=9); y -= 4*mm

            writeln("Opinia techniczno-finansowa — CWU (skrót)", bold=True); y -= 2*mm
            writeln("Wzory autorskie PROF INSTAL © 2025 Maciej Ślusarczyk:", bold=True, size=9)
            writeln("Q = m·c·ΔT;  Q[GJ/m³]=(1000·4,19·ΔT)/1e6;  koszt_teor = Q·cena_ciepła_brutto;  η = koszt_teor/rachunek")
            writeln(f"Rachunek: {L['bill']:.2f} zł/m³ | Ciepło brutto: {L['price_GJ_brutto']:.2f} zł/GJ | ΔT: {L['dT']:.0f}°C")
            writeln(f"Q_teor: {L['q_per_m3']:.5f} GJ/m³ → koszt_teor: {L['cost_theor']:.2f} zł/m³ | η: {L['eta']*100:.1f}%")
            writeln(f"Strata: {L['loss_per_m3']:.2f} zł/m³ | Budynek: {L['loss_build_m']:,.2f} zł/m-c; {L['loss_build_y']:,.2f} zł/rok"); y -= 2*mm
            writeln("Scenariusze modernizacji:", bold=True)
            writeln(f"70% → koszt {L['cost70']:.2f} zł/m³ | oszcz. {L['save70_m3']:.2f} zł/m³")
            writeln(f"80% → koszt {L['cost80']:.2f} zł/m³ | oszcz. {L['save80_m3']:.2f} zł/m³")
            y -= 2*mm
            writeln("Oprogramowanie: PROF INSTAL - Dla Mieszkańców © 2025 Maciej Ślusarczyk", size=8)

            c.showPage(); c.save(); open_file_crossplatform(path)
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się utworzyć opinii (PDF).\n{e}")

    # ===== Porady =====
    def _show_tips(self):
        messagebox.showinfo(
            "Porady oszczędności",
            "© 2025 PROF INSTAL Maciej Ślusarczyk\n\n"
            "• Temp. CWU 55–60°C; kontrola anty-Legionella.\n"
            "• Docieplenia przewodów rozprowadzających i cyrkulacyjnych.\n"
            "• Zawory termostatyczne/równoważenie cyrkulacji; bilans pionów.\n"
            "• Sterowanie pomp wg temp. powrotu/harmonogramu.\n"
            "• Przegląd wymienników i automatyki węzła; eliminacja przegrzewów.\n\n"
            "Porady chronione prawami autorskimi."
        )

    # === STOPKA ===
    def _create_copyright_footer(self, parent):
        sep = ttk.Frame(parent, height=2, style="Bg.TFrame")
        sep.pack(fill="x", pady=(10, 0))
        footer = ttk.Frame(parent, style="Bg.TFrame", padding=(15, 8))
        footer.pack(fill="x", side="bottom")
        footer.columnconfigure(0, weight=1)
        footer.columnconfigure(1, weight=0)
        footer.columnconfigure(2, weight=1)
        left = ttk.Frame(footer, style="Bg.TFrame"); left.grid(row=0, column=0, sticky="w")
        ttk.Label(left, text="© 2025 PROF INSTAL Maciej Ślusarczyk", style="Info.TLabel", font=(FONT_FAMILY, 9, "bold")).pack(side="left")
        ttk.Label(left, text=" • Wszelkie prawa zastrzeżone", style="Info.TLabel", font=(FONT_FAMILY, 8)).pack(side="left")
        center = ttk.Frame(footer, style="Bg.TFrame"); center.grid(row=0, column=1, padx=30)
        ttk.Label(center, text="Model obliczeniowy chroniony prawami autorskimi", style="Info.TLabel", font=(FONT_FAMILY, 8)).pack()
        right = ttk.Frame(footer, style="Bg.TFrame"); right.grid(row=0, column=2, sticky="e")
        ttk.Label(right, text="prof.instal@example.com", style="Info.TLabel", font=(FONT_FAMILY, 8)).pack(side="right")

if __name__ == "__main__":
    App().mainloop()
